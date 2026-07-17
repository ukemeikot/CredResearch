"""Render a document (ProseMirror/Tiptap JSON per section) to a .docx, applying the template's
format rules (FR-DOC-5/6). Returns the file bytes. PDF (Gotenberg) + ZIP bundle build on this."""

from __future__ import annotations

import io

from docx import Document as Docx
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt


def _add_inline(paragraph, node: dict) -> None:
    """Append a ProseMirror inline node (text with marks) to a docx paragraph."""
    if node.get("type") != "text":
        return
    run = paragraph.add_run(node.get("text", ""))
    for mark in node.get("marks", []) or []:
        t = mark.get("type")
        if t == "bold":
            run.bold = True
        elif t == "italic":
            run.italic = True
        elif t in ("underline", "u"):
            run.underline = True


def _block_text(doc, node: dict) -> None:
    """Render a ProseMirror block node into the docx."""
    t = node.get("type")
    content = node.get("content", []) or []
    if t == "heading":
        level = min(int(node.get("attrs", {}).get("level", 2)), 4)
        p = doc.add_heading(level=level)
        for c in content:
            _add_inline(p, c)
    elif t == "paragraph":
        p = doc.add_paragraph()
        for c in content:
            _add_inline(p, c)
    elif t == "blockquote":
        for c in content:
            p = doc.add_paragraph(style="Intense Quote")
            for cc in c.get("content", []) or []:
                _add_inline(p, cc)
    elif t in ("bulletList", "orderedList"):
        style = "List Bullet" if t == "bulletList" else "List Number"
        for item in content:  # listItem
            for block in item.get("content", []) or []:
                p = doc.add_paragraph(style=style)
                for cc in block.get("content", []) or []:
                    _add_inline(p, cc)
    else:
        # Fallback: render any nested content as a paragraph.
        p = doc.add_paragraph()
        for c in content:
            _add_inline(p, c)


def build_docx(title: str, sections: list[dict], format_rule: dict | None) -> bytes:
    doc = Docx()

    fr = format_rule or {}
    normal = doc.styles["Normal"]
    normal.font.name = fr.get("font_family", "Times New Roman")
    normal.font.size = Pt(float(fr.get("font_size_pt", 12)))
    spacing = float(fr.get("line_spacing", 2.0))
    normal.paragraph_format.line_spacing = spacing
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    if title:
        doc.add_heading(title, level=0)

    current_chapter = None
    for section in sections:
        chapter = section.get("chapter")
        if chapter and chapter != current_chapter:
            doc.add_heading(chapter, level=1)
            current_chapter = chapter
        heading = section.get("heading")
        if heading:
            doc.add_heading(heading, level=2)
        content = section.get("content") or {}
        for node in content.get("content", []) or []:
            _block_text(doc, node)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
